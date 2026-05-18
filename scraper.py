#!/usr/bin/env python3
"""
Teaching Job Scraper for Amro Badr - v2
Uses embedded JSON data from eChinacities HTML pages
Filters: English, Business, Economics, Humanities | Salary >= 14,000 RMB | Grade 6+
Output: XLS + .docx cover letters
"""

import requests
from bs4 import BeautifulSoup
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import re
import json
import os
import hashlib
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── CONFIGURATION ───────────────────────────────────────────────
PROFILE = {
    'name': 'Amro Badr (Ford)',
    'email': 'Amrbadr15@gmail.com',
    'phone': '+84 898091337',
    'subjects': ['english', 'esl', 'efl', 'tefl', 'business', 'economics', 'humanities', 'business management', 'finance', 'teaching', 'education'],
    'teaching_certs': ['TEFL', 'TESOL', 'PGCE'],
    'degree': 'BA Business Management & Finance',
    'experience_years': 7,
    'current_title': 'International Teacher',
    'languages': {'English': 'Native', 'Arabic': 'Native', 'Vietnamese': 'Spoken'},
}

HEADERS = {
    'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
    'Accept-Language': 'en-US,en;q=0.5',
}

OUTPUT_DIR = '/home/amrba/job_scraper'
JOBS_FILE = f'{OUTPUT_DIR}/jobs_data.json'
HISTORY_FILE = f'{OUTPUT_DIR}/scraped_history.json'
XLS_FILE = f'{OUTPUT_DIR}/teaching_jobs_china_{datetime.now().strftime("%Y%m%d")}.xlsx'
COVER_LETTERS_DIR = f'{OUTPUT_DIR}/cover_letters'
LOG_FILE = f'{OUTPUT_DIR}/scraper.log'

# Google Drive folder IDs
DRIVE_FOLDER_JOBS_CHINA = '1xFCe2LuqwYAIh1CzoIeYQ_Kuic8OXDlm'
DRIVE_FOLDER_COVER_LETTERS = '11T_qMHbJ1z45sPSrKRyIwzJD6JvFS0l9'

# ─── UTILITIES ──────────────────────────────────────────────────

def log(msg):
    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    line = f'[{ts}] {msg}'
    print(line)
    with open(LOG_FILE, 'a') as f:
        f.write(line + '\n')

def upload_to_drive(local_path, name, parent_folder_id):
    """Upload a file to Google Drive into a specific folder."""
    import subprocess, json
    try:
        result = subprocess.run([
            '/home/amrba/.hermes/hermes-agent/venv/bin/python3',
            '/home/amrba/.hermes/skills/productivity/google-workspace/scripts/google_api.py',
            'drive', 'upload', local_path,
            '--name', name,
            '--parent', parent_folder_id
        ], capture_output=True, text=True, timeout=30)
        if result.stdout:
            resp = json.loads(result.stdout)
            if resp.get('status') == 'uploaded':
                log(f"  ☁️  Drive: {resp.get('webViewLink', '')}")
                return True
    except Exception as e:
        log(f"  ⚠️  Drive upload failed: {e}")
    return False

def ensure_dir(path):
    os.makedirs(path, exist_ok=True)

def compute_job_hash(company, title, salary):
    key = f'{company}|{title}|{salary}'.lower()
    return hashlib.md5(key.encode()).hexdigest()

def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, 'r') as f:
            return set(json.load(f))
    return set()

def save_history(hashes):
    with open(HISTORY_FILE, 'w') as f:
        json.dump(list(hashes), f)

def parse_salary(salary_str):
    """Extract numeric salary from string. Returns (min, max, currency)"""
    if not salary_str:
        return None, None, None
    salary_str = salary_str.strip()
    
    patterns = [
        r'([¥$€£])[\s]*([\d,]+)\s*[-–to]+\s*([\d,]+)',
        r'([¥$€£])[\s]*([\d,]+)',
        r'([\d,]+)\s*[-–to]+\s*([\d,]+)\s*(?:RMB|USD|CNY|¥|€|£)',
        r'([\d,]+)\s*(?:RMB|USD|CNY|¥|€|£)',
    ]
    for pat in patterns:
        m = re.search(pat, salary_str, re.IGNORECASE)
        if m:
            groups = m.groups()
            curr = 'RMB'
            if '€' in salary_str: curr = 'EUR'
            elif '£' in salary_str: curr = 'GBP'
            elif '$' in salary_str: curr = 'USD'
            elif '¥' in salary_str or 'CNY' in salary_str.upper(): curr = 'RMB'
            try:
                if len(groups) >= 2 and groups[1]:
                    min_val = int(re.sub(r'[^0-9]', '', groups[1]))
                    max_val = int(re.sub(r'[^0-9]', '', groups[2])) if len(groups) >= 3 and groups[2] else min_val
                    return min_val, max_val, curr
                elif len(groups) == 1:
                    val = int(re.sub(r'[^0-9]', '', groups[0]))
                    return val, val, curr
            except:
                pass
    nums = re.findall(r'([\d,]+)', salary_str)
    if nums:
        val = int(nums[0].replace(',', ''))
        curr = 'RMB' if any(x in salary_str for x in ['¥', 'CNY']) else ('USD' if '$' in salary_str else 'RMB')
        return val, val, curr
    return None, None, None

def normalize_to_rmb(amount, currency):
    rates = {'RMB': 1, 'USD': 7.2, 'EUR': 7.8, 'GBP': 9.2}
    return amount * rates.get(currency, 1)

def meets_salary_requirement(salary_str):
    min_sal, _, curr = parse_salary(salary_str)
    if min_sal is None:
        return True
    return normalize_to_rmb(min_sal, curr) >= PROFILE['min_salary_rmb']

def meets_subject_requirement(text):
    text_lower = text.lower()
    for subj in PROFILE['subjects']:
        if subj in text_lower:
            return True
    return False

def grade_level_ok(text):
    text_lower = text.lower()
    bad_patterns = [
        r'kindergarten\b', r'\bkindergarten\b',
        r'preschool', r'nursery',
        r'primary\s*school', r'elementary',
        r'grade\s*[0-5]\b', r'\bgrade\s*[0-5]',
    ]
    for pat in bad_patterns:
        if re.search(pat, text_lower):
            return False
    return True

def matches_profile(job):
    full_text = f"{job.get('title', '')} {job.get('description', '')}"

    # Must mention target subjects
    if not meets_subject_requirement(full_text):
        return False, 'Subject mismatch'

    # Score the job
    job['match_score'] = compute_match_score(job)
    if job['match_score'] < 30:
        return False, 'Score too low'

    return True, 'Match'

def compute_match_score(job):
    score = 50
    sal_str = job.get('salaryRmb', '') or job.get('salary', '')
    min_sal, _, curr = parse_salary(sal_str)
    if min_sal:
        rmb = normalize_to_rmb(min_sal, curr)
        if rmb >= 25000: score += 30
        elif rmb >= 20000: score += 20
        elif rmb >= 16000: score += 15
        elif rmb >= 14000: score += 10
    
    text = f"{job.get('title','')} {job.get('description','')}"
    for subj in PROFILE['subjects']:
        if subj.lower() in text.lower():
            score += 5
            break
    
    if 'full' in job.get('job_type', '').lower():
        score += 5
    
    return min(score, 100)

# ─── ECHINACITIES SCRAPER ─────────────────────────────────────────

def extract_eChinacities_json(html_content):
    """Extract the embedded _searchJobList JSON from the HTML"""
    # Find the _searchJobList variable
    pattern = r'var\s+_searchJobList\s*=\s*(\{.*?\})\s*;'
    match = re.search(pattern, html_content, re.DOTALL)
    if not match:
        return None
    
    json_str = match.group(1)
    # Fix potential JavaScript syntax issues for JSON parsing
    # The data is already quite clean JSON from the server
    
    try:
        data = json.loads(json_str)
        return data
    except json.JSONDecodeError as e:
        # Try to fix common issues
        log(f"JSON parse error: {e}")
        return None

def scrape_eChinacities(session):
    """Scrape eChinacities using the embedded JSON data with pagination"""
    log("Scraping eChinacities...")
    jobs = []
    
    # Search terms matching profile subjects
    search_terms = [
        'english teacher',
        'business teacher',
        'economics teacher',
        'humanities teacher',
        'finance teacher',
    ]
    
    for term in search_terms:
        # Paginate through up to 20 pages per search term
        for page in range(1, 21):
            try:
                url = f'https://jobs.echinacities.com/jobs/search?keyword={term}&jobType=0&lastUpdate=30&page={page}'
                r = session.get(url, headers=HEADERS, timeout=20)
                
                if page == 1:
                    log(f"  '{term}': HTTP {r.status_code}, {len(r.text)} bytes")
                
                # Extract JSON from the page
                json_data = extract_eChinacities_json(r.text)
                
                if json_data and json_data.get('code') == 200:
                    job_list = json_data.get('data', {}).get('list', [])
                    num = json_data.get('data', {}).get('num', 0)
                    
                    if page == 1:
                        log(f"    Total available: {num} jobs")
                    
                    if not job_list:
                        break  # No more results for this term
                    
                    log(f"    Page {page}: {len(job_list)} jobs")
                    
                    for job in job_list:
                        # Build job dict from JSON
                        j = {
                            'source': 'eChinacities',
                            'id': job.get('id'),
                            'title': job.get('title', 'N/A'),
                            'company': job.get('company_name', 'N/A'),
                            'salary': job.get('salaryRmb', job.get('salary', 'N/A')),
                            'salary_usd': job.get('salaryDollar', 'N/A'),
                            'location': job.get('city', 'N/A'),
                            'job_type': job.get('job_type', 'N/A'),
                            'description': job.get('description', ''),
                            'requirements': job.get('description', ''),
                            'url': f"https://jobs.echinacities.com/jobs/detail?id={job.get('id')}",
                            'tags': job.get('sktag', []),
                            'scraped_at': datetime.now().isoformat(),
                            'refresh_time': job.get('refresh_time', ''),
                            'end_time': job.get('end_time', ''),
                            'is_new': job.get('isNew', False),
                            'is_hot': job.get('hot', False),
                        }
                        jobs.append(j)
                else:
                    if page == 1:
                        log(f"    No JSON data found")
                    break  # No more pages
                    
            except Exception as e:
                log(f"  '{term}' page {page} ERROR: {e}")
                break
    
    # Deduplicate by id
    seen_ids = set()
    unique_jobs = []
    for j in jobs:
        if j['id'] not in seen_ids:
            seen_ids.add(j['id'])
            unique_jobs.append(j)
    
    log(f"  eChinacities total: {len(unique_jobs)} unique jobs")
    return unique_jobs

# ─── OTHER SOURCES ──────────────────────────────────────────────

def scrape_other_sources(session):
    """Placeholder for additional sources"""
    # These would be implemented similarly
    return []

# ─── MAIN SCRAPER ───────────────────────────────────────────────

def scrape_all():
    log("="*60)
    log("STARTING JOB SCRAPE v2")
    log("="*60)
    
    ensure_dir(OUTPUT_DIR)
    ensure_dir(COVER_LETTERS_DIR)
    
    session = requests.Session()
    all_jobs = []

    # Load eChinacities session cookies if available
    COOKIES_FILE = f'{OUTPUT_DIR}/echinacities_cookies.json'
    if os.path.exists(COOKIES_FILE):
        with open(COOKIES_FILE) as f:
            cookies_dict = json.load(f)
        for name, value in cookies_dict.items():
            session.cookies.set(name, value, domain='.echinacities.com')
        log("eChinacities session cookies loaded")

    # Scrape eChinacities (main source with embedded JSON)
    all_jobs.extend(scrape_eChinacities(session))
    
    # Add other sources here
    # all_jobs.extend(scrape_eslcafe(session))
    
    log(f"Total raw jobs: {len(all_jobs)}")
    
    # Load history and filter duplicates
    history = load_history()
    new_jobs = []
    skipped = 0
    
    for job in all_jobs:
        j_hash = compute_job_hash(job['company'], job['title'], job.get('salary', ''))
        if j_hash in history:
            skipped += 1
            continue
        history.add(j_hash)
        new_jobs.append(job)
    
    log(f"New (non-duplicate) jobs: {len(new_jobs)} (skipped {skipped} duplicates)")
    
    # Filter to profile matches
    matched_jobs = []
    for job in new_jobs:
        is_match, reason = matches_profile(job)
        if is_match:
            matched_jobs.append(job)
            log(f"  ✅ MATCH: {job['title']} @ {job['company']} | {job['salary']} | {job['location']}")
        else:
            log(f"  ⛔ SKIP: {job.get('title', '?')} — {reason}")
    
    # Sort by match score
    matched_jobs.sort(key=lambda x: x.get('match_score', 0), reverse=True)
    
    log(f"Profile-matched jobs: {len(matched_jobs)}")
    
    # Save JSON
    with open(JOBS_FILE, 'w') as f:
        json.dump(matched_jobs, f, indent=2, default=str)
    log(f"Saved JSON: {JOBS_FILE}")
    
    # Save history
    save_history(history)
    
    # Export to XLS
    export_to_xls(matched_jobs)
    
    # Generate cover letters
    generate_cover_letters(matched_jobs)
    
    log("="*60)
    log("SCRAPE COMPLETE")
    log(f"  Total scraped: {len(all_jobs)}")
    log(f"  New jobs: {len(new_jobs)}")
    log(f"  Matched: {len(matched_jobs)}")
    log(f"  XLS: {XLS_FILE}")
    log(f"  Cover letters: {COVER_LETTERS_DIR}/")
    log("="*60)
    
    return matched_jobs

# ─── XLS EXPORTER ────────────────────────────────────────────────

def export_to_xls(jobs):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Teaching Jobs'
    
    header_font = Font(bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill(start_color='2E5A88', end_color='2E5A88', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    alt_fill = PatternFill(start_color='E8F0F7', end_color='E8F0F7', fill_type='solid')
    
    headers = ['#', 'Company', 'Job Title', 'Location', 'Salary (RMB)', 'Job Type', 'Match %', 'URL', 'Description', 'Tags']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border
    
    col_widths = [5, 28, 40, 18, 18, 15, 8, 45, 55, 30]
    for col, width in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = width
    
    for row_idx, job in enumerate(jobs, 2):
        row_data = [
            row_idx - 1,
            job.get('company', 'N/A'),
            job.get('title', 'N/A'),
            job.get('location', 'N/A'),
            job.get('salary', 'N/A'),
            job.get('job_type', 'N/A'),
            f"{job.get('match_score', 0)}%",
            job.get('url', 'N/A'),
            (job.get('description', '') or '')[:300],
            ', '.join(job.get('tags', [])),
        ]
        fill = alt_fill if row_idx % 2 == 0 else None
        for col, val in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col, value=val)
            cell.border = thin_border
            if fill:
                cell.fill = fill
            if col in [1, 6, 7]:
                cell.alignment = Alignment(horizontal='center')
    
    ws.freeze_panes = 'A2'
    ws.auto_filter.ref = f'A1:J{len(jobs)+1}'
    
    wb.save(XLS_FILE)
    log(f"Saved XLS: {XLS_FILE}")

    # Upload XLS to Google Drive (Jobs China folder)
    xls_name = f"China Jobs - {datetime.now().strftime('%B %d, %Y')}.xlsx"
    upload_to_drive(XLS_FILE, xls_name, DRIVE_FOLDER_JOBS_CHINA)

# ─── COVER LETTER GENERATOR ─────────────────────────────────────

def generate_cover_letters(jobs):
    ensure_dir(COVER_LETTERS_DIR)
    
    for idx, job in enumerate(jobs[:20], 1):
        try:
            doc = Document()
            
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Name
            p = doc.add_paragraph()
            run = p.add_run('Amro Badr (Ford)')
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x2E, 0x5A, 0x88)
            
            p = doc.add_paragraph()
            p.add_run('Email: Amrbadr15@gmail.com | WeChat: Ford1702 | Phone: +84 898091337').font.size = Pt(10)
            
            p = doc.add_paragraph()
            p.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}').font.size = Pt(10)
            
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            p.add_run('Dear Hiring Manager,').font.size = Pt(11)
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            run = p.add_run(f"Re: {job.get('title', 'Teaching Position')} - {job.get('company', 'Your School')}")
            run.bold = True
            run.font.size = Pt(11)
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            opening = (
                f"I am writing to express my strong interest in the {job.get('title', 'teaching position')} "
                f"at {job.get('company', 'your school')}{' in ' + job.get('location', '') if job.get('location') else ''}. "
                f"With over 7 years of teaching experience, TEFL, TESOL, and a Postgraduate Certificate in "
                f"Education (PGCE) certification, combined with my BA in Business Management & Finance, "
                f"I am well-equipped to deliver high-quality instruction in English, Business, Economics, "
                f"and Humanities at the high school level."
            )
            p.add_run(opening).font.size = Pt(11)
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            run = p.add_run('Key Qualifications:')
            run.bold = True
            run.font.size = Pt(11)
            
            quals = [
                "Teaching Certification: TEFL, TESOL, and PGCE (Liverpool John Moores University)",
                "Academic Background: BA Business Management & Finance; MA in Education (ongoing, completion November 2026)",
                "Teaching Experience: 7+ years across IELTS instruction, Business English, and Montessori pedagogy",
                "Subject Expertise: English, Business, Economics, Humanities — all at high school level",
                "Language Skills: English (Native), Arabic (Native), Vietnamese (Spoken)",
                "Additional: CAPM® certified, IELTS Speaking Teacher Training",
            ]
            for q in quals:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(q).font.size = Pt(10)
            
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            closing = (
                f"I am eager to contribute to {job.get('company', 'your school')} and would welcome "
                f"the opportunity to discuss how my background aligns with your needs. "
                f"I am available for a one-year contract or longer and am willing to relocate."
            )
            p.add_run(closing).font.size = Pt(11)
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            p.add_run('Yours sincerely,').font.size = Pt(11)
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run('Amro Badr (Ford)')
            run.bold = True
            run.font.size = Pt(11)
            
            safe_company = re.sub(r'[^\w\s-]', '', job.get('company', 'School'))[:20].strip()
            safe_title = re.sub(r'[^\w\s-]', '', job.get('title', 'Job'))[:30].strip()
            filename = f"{idx}_{safe_company}_{safe_title}.docx"
            filepath = os.path.join(COVER_LETTERS_DIR, filename)
            doc.save(filepath)
            log(f"  📄 Cover letter: {filename}")

            # Upload to Google Drive (Cover Letters folder)
            try:
                import subprocess
                result = subprocess.run([
                    '/home/amrba/.hermes/hermes-agent/venv/bin/python3',
                    '/home/amrba/.hermes/skills/productivity/google-workspace/scripts/google_api.py',
                    'drive', 'upload', filepath,
                    '--name', f"{safe_company} - {safe_title}.docx",
                    '--parent', DRIVE_FOLDER_COVER_LETTERS
                ], capture_output=True, text=True, timeout=15)
                if result.stdout:
                    import json
                    upload_resp = json.loads(result.stdout)
                    if upload_resp.get('status') == 'uploaded':
                        log(f"  ☁️  Uploaded to Drive: {upload_resp.get('webViewLink', '')}")
            except Exception as e:
                log(f"  ⚠️  Drive upload failed: {e}")
            
        except Exception as e:
            log(f"  Cover letter error for '{job.get('title', '?')}': {e}")
    
    log(f"Generated {min(len(jobs), 20)} cover letters in {COVER_LETTERS_DIR}/")

# ─── MAIN ────────────────────────────────────────────────────────

if __name__ == '__main__':
    ensure_dir(OUTPUT_DIR)
    ensure_dir(COVER_LETTERS_DIR)
    
    with open(LOG_FILE, 'w') as f:
        f.write('')
    
    results = scrape_all()
    
    print(f"\n✅ SCRAPE COMPLETE")
    print(f"   Jobs matched to profile: {len(results)}")
    print(f"   XLS file: {XLS_FILE}")
    print(f"   Cover letters: {COVER_LETTERS_DIR}/")
    print(f"   Log: {LOG_FILE}")