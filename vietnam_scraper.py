#!/usr/bin/env python3
"""
Vietnam Teaching Job Scraper for Amro Badr
Sources: LinkedIn (Playwright), Vietnamworks (Playwright), Glassdoor (requests)
Filters: Grade 6+, English/Business/Economics/Humanities, Full-time, >=50M VND/month
Output: XLS + .docx cover letters (max 20)
"""

import os, sys, re, json, hashlib
from datetime import datetime
import time, random

# ─── PATHS ────────────────────────────────────────────────────────────────────
BASE_DIR = '/home/amrba/job_scraper'
VENV_PYTHON = '/home/amrba/job_scraper/venv/bin/python3'
VIETNAM_DIR = f'{BASE_DIR}/vietnam'
COVER_LETTERS_DIR = f'{VIETNAM_DIR}/cover_letters'
os.makedirs(VIETNAM_DIR, exist_ok=True)
os.makedirs(COVER_LETTERS_DIR, exist_ok=True)

DATA_FILE = f'{VIETNAM_DIR}/jobs_data.json'
HISTORY_FILE = f'{VIETNAM_DIR}/scraped_history.json'
LOG_FILE = f'{VIETNAM_DIR}/scraper.log'

# ─── PLAYWRIGHT SETUP ─────────────────────────────────────────────────────────
PLAYWRIGHT_BROWSER_PATH = '/snap/chromium/3423/usr/lib/chromium-browser/chrome'
PLAYWRIGHT_AVAILABLE = os.path.exists(PLAYWRIGHT_BROWSER_PATH)

_playwright_instance = None
_browser = None

def _get_pw():
    global _playwright_instance, _browser
    if _playwright_instance is None:
        from playwright.sync_api import sync_playwright
        _playwright_instance = sync_playwright().start()
        _browser = _playwright_instance.chromium.launch(
            headless=True,
            executable_path=PLAYWRIGHT_BROWSER_PATH,
            args=['--no-sandbox', '--disable-dev-shm-usage', '--disable-gpu',
                  '--disable-software-rasterizer', '--disable-web-security']
        )
    return _playwright_instance, _browser

def _close_pw():
    global _playwright_instance, _browser
    if _browser:
        _browser.close()
        _browser = None
    if _playwright_instance:
        _playwright_instance.stop()
        _playwright_instance = None

# ─── PROFILE ──────────────────────────────────────────────────────────────────
PROFILE = {
    'name': 'Amro Badr (Ford)',
    'email': 'amrbadr15@gmail.com',
    'phone': '+84 898091337',
    'subjects': ['english', 'business', 'economics', 'humanities',
                 'business management', 'finance'],
    'min_salary_vnd': 50_000_000,   # 50M VND/month minimum
    'target_salary_vnd': 70_000_000, # 70M VND target
    'min_grade': 6,
    'require_fulltime': True,
}

SEARCH_TERMS = [
    'english teacher vietnam',
    'ielts teacher vietnam',
    'business teacher vietnam',
    'economics teacher vietnam',
    'esl teacher vietnam',
]

# ─── LOGGING ──────────────────────────────────────────────────────────────────
def log(msg):
    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    line = f'[{ts}] {msg}'
    print(line)
    with open(LOG_FILE, 'a') as f:
        f.write(line + '\n')

# ─── SALARY PARSING ────────────────────────────────────────────────────────────
# Convert everything to VND/month for consistent comparison
USD_TO_VND = 25_000
HOURS_PER_MONTH = 160  # full-time teaching hours

def parse_salary_vnd(text):
    """Returns (min_vnd_month, max_vnd_month) normalized to monthly VND."""
    if not text:
        return (None, None)
    text = text.lower().strip()

    # ── Per-hour patterns ────────────────────────────────────────────────────
    hourly_patterns = [
        r'([\d,]+)\s*(?:k|vnd|₫)\s*(?:per hour|/hour|hour|hrs?)/?(?:month)?',
        r'([\d,]+)\s*(?:vnd|₫)\s*(?:per hour|/hour)',
        r'([\d,]+)\s*k\s*(?:per hour|/hour|hour)',
        r'up to\s+([\d,]+)\s*(?:k|vnd|₫)\s*(?:per hour|/hour)',
    ]
    for pat in hourly_patterns:
        m = re.search(pat, text)
        if m:
            raw = m.group(1).replace(',', '').strip()
            if not raw:
                return (None, None)
            val = float(raw) * 1_000
            monthly = val * HOURS_PER_MONTH
            return (monthly, monthly)

    # ── Per-month patterns (triệu / million) ────────────────────────────────
    month_patterns = [
        r'([\d,]+)\s*(?:triệu|tr)\s*(?:vnd|₫)?',
        r'([\d,]+)\s*(?:million|m)\s*(?:vnd|₫)?',
        r'([\d,]+)\s*(?:vnd|₫)\s*(?:per month|/month)?',
        r'([\d,]+)\s*-\s*([\d,]+)\s*(?:triệu|tr|million)',
        r'([\d,]+)\s*-\s*([\d,]+)\s*(?:vnd|₫)',
    ]
    for pat in month_patterns:
        m = re.search(pat, text)
        if m:
            raw1 = m.group(1).replace(',', '').strip()
            if not raw1:
                return (None, None)
            if len(m.groups()) == 2:
                raw2 = m.group(2).replace(',', '').strip()
                if not raw2:
                    return (None, None)
                lo = float(raw1) * 1_000_000
                hi = float(raw2) * 1_000_000
            else:
                val = float(raw1) * 1_000_000
                lo = hi = val
            return (lo, hi)

    # ── USD patterns ─────────────────────────────────────────────────────────
    usd_range = r'\$([\d,]+)\s*-\s*\$?([\d,]+)\s*(?:usd|usd/month|usd/yr)'
    m = re.search(usd_range, text)
    if m:
        raw1 = m.group(1).replace(',', '').strip()
        raw2 = m.group(2).replace(',', '').strip()
        if not raw1 or not raw2:
            return (None, None)
        lo = float(raw1) * USD_TO_VND
        hi = float(raw2) * USD_TO_VND
        return (lo, hi)
    usd_single = r'\$([\d,]+)\s*(?:usd|usd/month|usd/yr)'
    m = re.search(usd_single, text)
    if m:
        raw = m.group(1).replace(',', '').strip()
        if not raw:
            return (None, None)
        val = float(raw) * USD_TO_VND
        return (val, val)

    return (None, None)

# ─── LINKEDIN SCRAPING ─────────────────────────────────────────────────────────
def scrape_linkedin():
    log("Scraping LinkedIn Vietnam (Playwright)...")
    if not PLAYWRIGHT_AVAILABLE:
        log("  Playwright browser not found, skipping LinkedIn")
        return []

    _, browser = _get_pw()
    jobs = []
    seen_ids = set()

    for term in SEARCH_TERMS:
        encoded_term = term.replace(' ', '%20')
        search_url = (f'https://www.linkedin.com/jobs/search?keywords={encoded_term}'
                      f'&location=Vietnam&position=1&pageNum=0')
        log(f"  LinkedIn search: {term}")

        page = None
        try:
            page = browser.new_page()
            page.goto(search_url, timeout=30000, wait_until='domcontentloaded')
            page.wait_for_timeout(2000)

            html = page.content()
            job_ids = re.findall(r'urn:li:jobPosting:(\d+)', html)
            new_ids = [jid for jid in job_ids if jid not in seen_ids]
            seen_ids.update(job_ids)
            log(f"    Found {len(job_ids)} IDs total, {len(new_ids)} new")
            page.close()
            page = None

            if not new_ids:
                time.sleep(random.uniform(1, 2))
                continue

            for jid in new_ids:
                job = _fetch_linkedin_job(browser, jid)
                if job:
                    jobs.append(job)
                    log(f"    ✓ Job {job['id']}: {job['title'][:50]} @ {job['company']}")
                time.sleep(random.uniform(0.5, 1.0))

        except Exception as e:
            log(f"  Search error for '{term}': {e}")
        finally:
            if page:
                page.close()

        time.sleep(random.uniform(1, 2))

    log(f"  LinkedIn: collected {len(jobs)} raw jobs")
    return jobs

def _fetch_linkedin_job(browser, jid):
    """Fetch a single LinkedIn job detail page."""
    job_url = f'https://www.linkedin.com/jobs/view/{jid}'
    page = browser.new_page()
    try:
        page.goto(job_url, timeout=25000, wait_until='domcontentloaded')
        page.wait_for_timeout(1500)

        body_text = page.inner_text('body')
        html = page.content()

        if len(body_text) < 80:
            return None

        # Title — h1 or og:title
        title = None
        if page.locator('h1').count() > 0:
            h1 = page.locator('h1').first.text_content()
            if h1 and len(h1.strip()) > 3:
                title = h1.strip()

        # Company from page title tag
        company = None
        title_tag = re.search(r'<title>([^<]+)</title>', html)
        if title_tag:
            ts = title_tag.group(1)
            if ' hiring ' in ts:
                company = ts.split(' hiring ')[0].strip()
            elif ' | LinkedIn' in ts:
                company = ts.replace(' | LinkedIn', '').strip()

        # Location — look for Vietnamese cities in body
        location = 'Vietnam'
        loc_cities = ['ho chi minh', 'hanoi', 'da nang', 'hue', 'can tho',
                       'hai phong', 'nha trang', 'vung tau']
        for line in body_text.split('\n'):
            low = line.lower().strip()
            if any(c in low for c in loc_cities) and 3 < len(line.strip()) < 70:
                location = line.strip()
                break

        # Salary
        salary = None
        salary_patterns = [
            r'[\d,]+[\s]*(?:VND|₫|USD|\$)[\s]*(?:per month|/month|/hr|hour)?',
            r'[\d,]+[\s]*(?:triệu|tr|million|m)[\s]*(?:vnd|₫)?[\s]*(?:per month|/month)?',
            r'(?:VND|USD)[\s]*[\d,]+[\s]*-[\s]*[\d,]+',
            r'\$[\d,]+[\s]*-[\s]*\$[\d,]+',
            r'up to[\s\$]*[\d,]+',
        ]
        for pat in salary_patterns:
            m = re.search(pat, body_text, re.IGNORECASE)
            if m:
                salary = m.group(0).strip()
                break

        # Employment type
        job_type = 'Full-time'
        et_lower = body_text.lower()
        if re.search(r'part[\s-]?time', et_lower):
            job_type = 'Part-time'
        elif re.search(r'contract|seasonal|temp', et_lower):
            job_type = 'Contract'

        page.close()
        time.sleep(random.uniform(0.3, 0.7))

        if title or company:
            return {
                'source': 'LinkedIn',
                'id': jid,
                'title': (title or 'Unknown').replace('&amp;', '&'),
                'company': (company or 'Unknown').replace('&amp;', '&'),
                'salary': salary or '',
                'location': location,
                'job_type': job_type,
                'description': body_text[:4000],
                'requirements': body_text[:4000],
                'url': job_url,
                'tags': [],
                'scraped_at': datetime.now().isoformat(),
                'match_score': 0,
                'match_reasons': '',
            }
        return None

    except Exception:
        page.close()
        return None

# ─── VIETNAMWORKS SCRAPING ─────────────────────────────────────────────────────
def scrape_vietnamworks():
    log("Scraping Vietnamworks (Playwright)...")
    if not PLAYWRIGHT_AVAILABLE:
        log("  Playwright not available, skipping Vietnamworks")
        return []

    _, browser = _get_pw()
    jobs = []
    seen_urls = set()

    # Search pages for English teaching jobs in HCMC and Hanoi
    search_queries = [
        ('english teacher', 'ho-chi-minh-vn'),
        ('english teacher', 'ha-noi-vn'),
        ('ielts teacher', 'ho-chi-minh-vn'),
        ('business english', 'ho-chi-minh-vn'),
    ]

    for query, city in search_queries:
        for page_num in range(2):
            q_encoded = query.replace(' ', '-')
            url = (f'https://www.vietnamworks.com/en/jobs/search/{q_encoded}'
                   f',in-{city}?page={page_num}')
            log(f"    Vietnamworks: {query} {city} page {page_num}")

            page = None
            try:
                page = browser.new_page()
                page.goto(url, timeout=30000, wait_until='domcontentloaded')
                page.wait_for_timeout(2500)
                html = page.content()

                # Find job listing cards
                job_cards = page.locator('.job-item, .job-search-result-item, '
                                          '[data-job-id], .jobs-card').all()
                log(f"      Found {len(job_cards)} job cards on page")

                # Also try regex for job URLs
                job_urls = re.findall(r'href="(/en/jobs/view/\d+[^"]*)"', html)
                job_urls = list(set(job_urls))
                log(f"      Found {len(job_urls)} job URLs via regex")

                page.close()
                page = None

                for jurl in job_urls:
                    if jurl in seen_urls:
                        continue
                    seen_urls.add(jurl)
                    full_url = f'https://www.vietnamworks.com{jurl}'
                    job = _fetch_vietnamworks_job(browser, full_url)
                    if job:
                        jobs.append(job)
                        log(f"      ✓ {job['title'][:50]} @ {job['company']}")
                    time.sleep(random.uniform(0.5, 1.0))

            except Exception as e:
                log(f"      Page error: {e}")
            finally:
                if page:
                    page.close()

            time.sleep(random.uniform(1, 2))

    log(f"  Vietnamworks: collected {len(jobs)} raw jobs")
    return jobs

def _fetch_vietnamworks_job(browser, url):
    """Fetch a single Vietnamworks job detail page."""
    page = browser.new_page()
    try:
        page.goto(url, timeout=30000, wait_until='domcontentloaded')
        page.wait_for_timeout(2000)

        body_text = page.inner_text('body')
        if len(body_text) < 80:
            return None

        # Try to extract title
        title = None
        title_selectors = ['h1.job-title', 'h1.title', '.job-header h1',
                            'h1[itemprop="title"]', '.top-card-info h1']
        for sel in title_selectors:
            if page.locator(sel).count() > 0:
                t = page.locator(sel).first.text_content()
                if t and len(t.strip()) > 3:
                    title = t.strip()
                    break

        if not title and page.locator('h1').count() > 0:
            title = page.locator('h1').first.text_content().strip()

        # Company
        company = None
        company_selectors = ['.company-name', '.employer-name', '[itemprop="hiringOrganization"]',
                             '.company-logo-wrap img']
        for sel in company_selectors:
            if page.locator(sel).count() > 0:
                c = page.locator(sel).first.text_content()
                if c and len(c.strip()) > 1:
                    company = c.strip()
                    break

        if not company:
            c_match = re.search(r'(?:company|employer)[:\-]?\s*([^\n<]{3,60})', body_text, re.I)
            if c_match:
                company = c_match.group(1).strip()

        # Location
        location = 'Vietnam'
        loc_patterns = ['ho chi minh', 'hanoi', 'da nang', 'can tho', 'hai phong']
        for line in body_text.split('\n'):
            low = line.lower().strip()
            if any(c in low for c in loc_patterns) and 3 < len(line.strip()) < 70:
                location = line.strip()
                break

        # Salary
        salary = None
        sal_patterns = [
            r'([\d,]+)\s*(?:triệu|tr)\s*(?:vnd|₫)?(?:-\s*([\d,]+)\s*(?:triệu|tr))?',
            r'([\d,]+)\s*(?:vnd|₫)(?:\s*-\s*([\d,]+)\s*(?:vnd|₫))?',
            r'\$\s*([\d,]+)\s*-\s*\$\s*([\d,]+)',
            r'[\d,]+[\s]*k[\s]*(?:per hour|/hour)?',
        ]
        for pat in sal_patterns:
            m = re.search(pat, body_text, re.I)
            if m:
                salary = m.group(0).strip()
                break

        # Job type
        job_type = 'Full-time'
        if re.search(r'part[\s-]?time', body_text.lower()):
            job_type = 'Part-time'
        elif re.search(r'contract', body_text.lower()):
            job_type = 'Contract'

        page.close()
        time.sleep(random.uniform(0.3, 0.7))

        if title:
            return {
                'source': 'Vietnamworks',
                'id': re.search(r'/jobs/view/(\d+)', url).group(1) if re.search(r'/jobs/view/(\d+)', url) else url,
                'title': title.replace('&amp;', '&'),
                'company': (company or 'Unknown').replace('&amp;', '&'),
                'salary': salary or '',
                'location': location,
                'job_type': job_type,
                'description': body_text[:4000],
                'requirements': body_text[:4000],
                'url': url,
                'tags': [],
                'scraped_at': datetime.now().isoformat(),
                'match_score': 0,
                'match_reasons': '',
            }
        return None

    except Exception:
        page.close()
        return None

# ─── PROFILE FILTER & SCORING ──────────────────────────────────────────────────
GRADE_KEYWORDS = {
    6:  ['grade 6', 'grade 6th', 'grade vi', 'lớp 6'],
    7:  ['grade 7', 'grade 7th', 'grade vii', 'lớp 7'],
    8:  ['grade 8', 'grade 8th', 'grade viii', 'lớp 8'],
    9:  ['grade 9', 'grade 9th', 'grade ix', 'lớp 9'],
    10: ['grade 10', 'grade 10th', 'grade x', 'lớp 10'],
    11: ['grade 11', 'grade 11th', 'grade xi', 'lớp 11'],
    12: ['grade 12', 'grade 12th', 'grade xii', 'lớp 12'],
    'high': ['high school', 'upper secondary', 'secondary school',
             'lớp 10', 'lớp 11', 'lớp 12', 'grades 10-12'],
    'uni': ['university', 'college', 'higher education'],
}
SUBJECT_KEYWORDS = [
    'english', 'esl', 'efl', 'business', 'economics', 'humanities',
    'social studies', 'global perspectives', 'finance', 'accounting',
    'ict', 'computing', 'cambridge', 'ib', 'a-level', 'ibdp',
    'academic writing', 'business english', 'ielts', 'toefl',
]
REJECT_KEYWORDS = [
    'kindergarten', 'pre-school', 'pre school', 'primary school',
    'elementary', 'grade 1', 'grade 2', 'grade 3', 'grade 4', 'grade 5',
    'lớp 1', 'lớp 2', 'lớp 3', 'lớp 4', 'lớp 5', 'mẫu giáo',
    'kids', 'children',
]

def grade_level(text):
    text = text.lower()
    if any(kw in text for kw in REJECT_KEYWORDS):
        return 0
    for grade in sorted(GRADE_KEYWORDS.keys(),
                        key=lambda x: (isinstance(x, str), x)):
        if any(kw in text for kw in GRADE_KEYWORDS[grade]):
            return grade if isinstance(grade, int) else (
                'high' if grade == 'high' else 'uni')
    return None

def subject_match(text):
    return any(kw in text.lower() for kw in SUBJECT_KEYWORDS)

def score_job(job):
    score = 0
    reasons = []
    full_text = (job.get('title', '') + ' ' +
                 job.get('description', '') + ' ' +
                 job.get('requirements', '')).lower()
    salary_text = job.get('salary', '') or ''

    # ── Grade scoring ──────────────────────────────────────────────────────
    grade = grade_level(full_text)
    if isinstance(grade, int) and grade >= PROFILE['min_grade']:
        score += 40
        reasons.append(f'Grade {grade}')
    elif grade == 'high':
        score += 30
        reasons.append('High School')
    elif grade == 'uni':
        score += 20
        reasons.append('University')
    elif grade is None:
        score += 10  # unspecified grade — don't penalise heavily
        reasons.append('Grade unspecified')
    else:
        reasons.append(f'Grade {grade} < min')

    # ── Salary scoring ─────────────────────────────────────────────────────
    lo, hi = parse_salary_vnd(salary_text)
    salary_display = ''
    if hi:
        salary_display = f'{hi/1e6:.0f}M'
    if hi and hi >= PROFILE['target_salary_vnd']:
        score += 30
        reasons.append(f'Salary {salary_display}M VND (target)')
    elif hi and hi >= PROFILE['min_salary_vnd']:
        score += 20
        reasons.append(f'Salary {salary_display}M VND (min met)')
    elif hi is None and salary_text:
        score += 5  # salary stated but couldn't parse
        reasons.append('Salary stated')
    elif hi and hi < PROFILE['min_salary_vnd']:
        reasons.append(f'Salary {salary_display}M VND (low)')

    # ── Subject match ──────────────────────────────────────────────────────
    if subject_match(full_text):
        score += 20
        reasons.append('Subject match')

    # ── Job type ───────────────────────────────────────────────────────────
    jt = job.get('job_type', 'Full-time')
    if jt == 'Full-time':
        score += 10
        reasons.append('Full-time')
    elif PROFILE['require_fulltime'] and jt != 'Full-time':
        score -= 20  # penalise part-time/contract

    job['match_score'] = score
    job['match_reasons'] = ', '.join(reasons)
    return score >= 70

def filter_and_score(jobs):
    scored = [j for j in jobs if score_job(j)]
    scored.sort(key=lambda j: j['match_score'], reverse=True)
    log(f"  Jobs matching profile (score >= 70): {len(scored)}")
    for j in scored:
        log(f"    Score {j['match_score']}: {j['title'][:55]} @ {j['company'][:30]} "
            f"| {j['job_type']} | {j['salary']}")
    return scored

# ─── DEDUP ─────────────────────────────────────────────────────────────────────
def job_hash(job):
    key = f"{job.get('company','')}|{job.get('title','')}|{job.get('salary','')}"
    return hashlib.md5(key.encode()).hexdigest()

def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE) as f:
            return set(json.load(f))
    return set()

def save_history(hashes):
    with open(HISTORY_FILE, 'w') as f:
        json.dump(sorted(hashes), f, indent=2)

def load_jobs_data():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE) as f:
            return json.load(f)
    return []

def save_jobs_data(jobs):
    with open(DATA_FILE, 'w') as f:
        json.dump(jobs, f, indent=2, ensure_ascii=False)

# ─── XLS EXPORT ────────────────────────────────────────────────────────────────
def export_xls(jobs, date_str):
    if not jobs:
        log("No jobs to export.")
        return
    xls_path = f'{VIETNAM_DIR}/vietnam_jobs_{date_str}.xlsx'
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Jobs'

    header_font = Font(bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill('solid', fgColor='003699')
    alt_fill = PatternFill('solid', fgColor='EEF2FF')
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    cols = ['Company Name', 'Job Title', 'Location', 'Salary (VND)',
            'Requirements', 'Score', 'Match Reasons', 'Source', 'URL', 'Job Type']
    ws.append(cols)
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

    for i, job in enumerate(jobs):
        row = [
            job.get('company', ''),
            job.get('title', ''),
            job.get('location', ''),
            job.get('salary', ''),
            (job.get('requirements', '') or job.get('description', ''))[:400],
            job.get('match_score', 0),
            job.get('match_reasons', ''),
            job.get('source', ''),
            job.get('url', ''),
            job.get('job_type', ''),
        ]
        ws.append(row)
        row_idx = i + 2
        if i % 2 == 1:
            for col in range(1, len(cols) + 1):
                ws.cell(row=row_idx, column=col).fill = alt_fill
        ws.cell(row=row_idx, column=6).alignment = Alignment(horizontal='center')
        for col in range(1, len(cols) + 1):
            ws.cell(row=row_idx, column=col).border = border

    ws.column_dimensions['A'].width = 28
    ws.column_dimensions['B'].width = 38
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 20
    ws.column_dimensions['E'].width = 55
    ws.column_dimensions['F'].width = 8
    ws.column_dimensions['G'].width = 32
    ws.column_dimensions['H'].width = 14
    ws.column_dimensions['I'].width = 45
    ws.column_dimensions['J'].width = 12
    ws.freeze_panes = 'A2'
    try:
        ws.auto_filter.ref = ws.dimensions
    except Exception:
        pass
    wb.save(xls_path)
    log(f"  XLS saved: {xls_path} ({len(jobs)} jobs)")

# ─── COVER LETTER ─────────────────────────────────────────────────────────────
COVER_TEMPLATE = """{name}
{email} | {phone}

Dear Hiring Manager,

I am writing to express my strong interest in the {title} position at {company}, as advertised on {source}. With seven years of dedicated English and test-preparation teaching experience — including five years as an IELTS instructor — combined with my TEFL/TESOL certification, PGCE studies through Liverpool John Moores University, and CAPM® and DASM credentials, I am confident I would be a valuable addition to your academic team.

In my current role as IELTS Teacher at RMIT English School (2018–2024), I have consistently helped students achieve band scores of 7.0 and above through structured, evidence-based instruction. My background spans IELTS, Business English, and Academic Writing, and I have taught across diverse learner levels from intermediate to advanced. This experience has equipped me with deep knowledge of curriculum design, differentiated instruction, and student assessment — skills directly applicable to {title} roles.

My PGCE studies (Liverpool John Moores University, 2024–2026) are strengthening my pedagogical practice with a particular focus on secondary and upper-secondary education. I hold a Bachelor's degree in Business Management & Finance and bring professional certifications in Project Management (CAPM®) and Agile Software Development (DASM), making me especially effective in Business, Economics, and Humanities classrooms where real-world application and cross-disciplinary thinking are valued.

I am experienced teaching Cambridge, IB, and national curricula, and I am fully conversant with modern educational technology platforms. I am eager to contribute to {company}'s academic mission and would welcome the opportunity to discuss how my background aligns with your team's needs.

Thank you for your consideration. I look forward to hearing from you.

Sincerely,
{name}
{email} | {phone}
"""

def generate_cover_letter(job):
    date_str = datetime.now().strftime('%Y%m%d')
    company_safe = re.sub(r'[^\w\s\-]', '', job.get('company', 'Company'))[:35]
    title_safe = re.sub(r'[^\w\s\-]', '', job.get('title', 'Position'))[:35]
    filename = f'{company_safe}_{title_safe}_{date_str}.docx'
    filepath = os.path.join(COVER_LETTERS_DIR, filename)

    # Detect grade from job text
    text = (job.get('title', '') + ' ' + job.get('description', '')).lower()
    grade = 10
    for g in range(6, 13):
        if any(kw in text for kw in GRADE_KEYWORDS.get(g, [])):
            grade = g
            break

    content = COVER_TEMPLATE.format(
        name=PROFILE['name'],
        email=PROFILE['email'],
        phone=PROFILE['phone'],
        title=job.get('title', 'the position'),
        company=job.get('company', 'your school'),
        source=job.get('source', 'LinkedIn'),
        grade=grade,
    )

    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    for para_text in content.split('\n\n'):
        p = doc.add_paragraph()
        p.add_run(para_text.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(8)
    doc.save(filepath)
    return filepath

# ─── MAIN ──────────────────────────────────────────────────────────────────────
def main():
    date_str = datetime.now().strftime('%Y%m%d')
    log(f"")
    log(f"=== Vietnam Job Scraper START: {date_str} ===")

    # Load existing data
    all_jobs = load_jobs_data()
    existing_hashes = load_history()
    log(f"  Loaded {len(all_jobs)} existing jobs, {len(existing_hashes)} dedup hashes")

    # ── Scrape ────────────────────────────────────────────────────────────────
    raw_jobs = []
    li_jobs = scrape_linkedin()
    raw_jobs.extend(li_jobs)

    vw_jobs = scrape_vietnamworks()
    raw_jobs.extend(vw_jobs)
    log(f"  Total raw jobs collected: {len(raw_jobs)}")

    if not raw_jobs:
        log("  No raw jobs collected — check Playwright/browser availability.")
        log("  Will re-export existing data if available.")
        _close_pw()
        existing = load_jobs_data()
        if existing:
            export_xls(sorted(existing, key=lambda x: x.get('match_score', 0), reverse=True), date_str)
        return

    # ── Score & filter ─────────────────────────────────────────────────────────
    matched = filter_and_score(raw_jobs)

    # ── Dedup against history ─────────────────────────────────────────────────
    new_jobs = []
    new_hashes = set(existing_hashes)
    for job in matched:
        h = job_hash(job)
        if h not in new_hashes:
            new_hashes.add(h)
            new_jobs.append(job)

    log(f"  New unique matched jobs: {len(new_jobs)}")

    # Update all_jobs — prepend new jobs to keep newest first
    all_jobs = new_jobs + all_jobs
    save_jobs_data(all_jobs)
    save_history(new_hashes)
    log(f"  Total jobs in data store: {len(all_jobs)}")

    # ── Export XLS ─────────────────────────────────────────────────────────────
    # Export all matched jobs (new + previously matched)
    exportable = [j for j in all_jobs if j.get('match_score', 0) >= 70]
    exportable.sort(key=lambda x: x.get('match_score', 0), reverse=True)
    export_xls(exportable, date_str)

    # ── Generate cover letters for new matches ────────────────────────────────
    cover_jobs = new_jobs[:20]
    for job in cover_jobs:
        try:
            path = generate_cover_letter(job)
            log(f"  Cover letter generated: {os.path.basename(path)}")
        except Exception as e:
            log(f"  Cover letter error for {job.get('title','?')}: {e}")

    log(f"  Cover letters generated: {len(cover_jobs)}")
    log(f"=== Vietnam Job Scraper END: {len(new_jobs)} new, "
        f"{len(cover_jobs)} cover letters ===")

    _close_pw()

if __name__ == '__main__':
    main()
