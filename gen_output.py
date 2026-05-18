#!/usr/bin/env python3
import json, re, os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

PROFILE = {'name':'Amro Badr (Ford)','email':'amrbadr15@gmail.com','phone':'+84 898091337'}
VIETNAM_DIR = '/home/amrba/job_scraper/vietnam'
COVER_LETTERS_DIR = VIETNAM_DIR + '/cover_letters'
os.makedirs(COVER_LETTERS_DIR, exist_ok=True)

GRADE_KEYWORDS = {
    6: ['grade 6','grade 6th','grade vi'], 7: ['grade 7','grade 7th','grade vii'],
    8: ['grade 8','grade 8th','grade viii'], 9: ['grade 9','grade 9th','grade ix'],
    10: ['grade 10','grade 10th','grade x'], 11: ['grade 11','grade 11th','grade xi'],
    12: ['grade 12','grade 12th','grade xii'],
    'high': ['high school','upper secondary','secondary school'],
}

with open(VIETNAM_DIR + '/jobs_data.json') as f:
    jobs = json.load(f)
print(f'Loaded {len(jobs)} jobs')
jobs.sort(key=lambda j: j.get('match_score',0), reverse=True)

xls_path = VIETNAM_DIR + '/vietnam_jobs_' + datetime.now().strftime('%Y%m%d') + '.xlsx'
wb = openpyxl.Workbook()
ws = wb.active
ws.title = 'Jobs'
cols = ['Company Name','Job Title','Location','Salary','Requirements','Match Score','Match Reasons','Source','URL']
ws.append(cols)
for cell in ws[1]:
    cell.font = Font(bold=True, color='FFFFFF', size=11)
    cell.fill = PatternFill('solid', fgColor='003699')
    cell.alignment = Alignment(horizontal='center')
for i, job in enumerate(jobs):
    row = [job.get('company',''), job.get('title',''), job.get('location',''), job.get('salary',''),
           (job.get('requirements','') or job.get('description',''))[:300],
           job.get('match_score',0), job.get('match_reasons',''), job.get('source',''), job.get('url','')]
    ws.append(row)
    if i % 2 == 1:
        for col in range(1, len(cols)+1):
            ws.cell(row=i+2, column=col).fill = PatternFill('solid', fgColor='EEF2FF')
    ws.cell(row=i+2, column=6).alignment = Alignment(horizontal='center')
ws.column_dimensions['A'].width = 25; ws.column_dimensions['B'].width = 35
ws.column_dimensions['C'].width = 18; ws.column_dimensions['D'].width = 22
ws.column_dimensions['E'].width = 50; ws.column_dimensions['F'].width = 12
ws.column_dimensions['G'].width = 30; ws.column_dimensions['H'].width = 10
ws.column_dimensions['I'].width = 30
ws.freeze_panes = 'A2'
ws.auto_filter.ref = 'A1:I' + str(len(jobs)+1)
wb.save(xls_path)
print(f'XLS saved: {xls_path}')

COVER = """{name}
{email} | {phone}

Dear Hiring Manager,

I am writing to express my strong interest in the {title} position at {company}, as advertised on {source}. With seven years of dedicated English and test-preparation teaching experience -- including five years as an IELTS instructor -- combined with my TEFL/TESOL certification, PGCE studies through Liverpool John Moores University, and CAPM(R) and DASM credentials, I am confident I would be a valuable addition to your academic team.

In my current role as IELTS Teacher at RMIT English School (2018-2024), I have consistently helped students achieve band scores of 7.0 and above through structured, evidence-based instruction. My background spans IELTS, Business English, and Academic Writing, and I have taught across diverse learner levels from intermediate to advanced. This experience has equipped me with deep knowledge of curriculum design, differentiated instruction, and student assessment.

My PGCE studies (Liverpool John Moores University, 2024-2026) are strengthening my pedagogical practice with a particular focus on secondary and upper-secondary education. I hold a Bachelor's degree in Business Management & Finance and bring professional certifications in Project Management (CAPM(R)) and Agile Software Development (DASM), making me especially effective in Business, Economics, and Humanities classrooms.

I am experienced teaching Cambridge, IB, and national curricula, and I am fully conversant with modern educational technology platforms. I am eager to contribute to {company}'s academic mission and would welcome the opportunity to discuss how my background aligns with your team's needs.

Thank you for your consideration. I look forward to hearing from you.

Sincerely,
{name}
{email} | {phone}
"""

count = 0
for job in jobs[:20]:
    if job.get('match_score',0) < 70:
        break
    cs = re.sub(r'[^\w\s\-]', '', job.get('company','Company'))[:40]
    ts = re.sub(r'[^\w\s\-]', '', job.get('title','Position'))[:40]
    fname = f'{cs}_{ts}_{datetime.now().strftime("%Y%m%d")}.docx'
    fpath = os.path.join(COVER_LETTERS_DIR, fname)
    g = 10
    txt = (job.get('title','') + ' ' + job.get('description','')).lower()
    for grade in range(6, 13):
        if any(kw in txt for kw in GRADE_KEYWORDS.get(grade, [])):
            g = grade; break
    content = COVER.format(name=PROFILE['name'], email=PROFILE['email'], phone=PROFILE['phone'],
        title=job.get('title','the position'), company=job.get('company','your school'),
        source=job.get('source','LinkedIn'), grade=g)
    doc = Document()
    for para in content.split('\n\n'):
        p = doc.add_paragraph()
        p.add_run(para.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(8)
    doc.save(fpath)
    count += 1
    print(f'Cover: {fname}')
print(f'Done: XLS + {count} cover letters')
