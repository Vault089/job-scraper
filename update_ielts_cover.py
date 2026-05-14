#!/usr/bin/env python3
"""Update Atlantic IELTS cover letter for 2026-2027 cycle."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

PROFILE = {'name': 'Amro Badr (Ford)', 'email': 'amrbadr15@gmail.com', 'phone': '+84 898091337'}

# Updated cover letter content - tailored for Atlantic Five-Star IELTS 2026-2027
# 730K VND/hr, 63 applicants, Sep 2026 start, students aged 15-18

content = f"""{PROFILE['name']}
{PROFILE['email']} | {PROFILE['phone']}

Dear Hiring Manager,

I am writing to express my strong interest in the Full-Time IELTS Teacher position at Atlantic Five-Star English, as part of the 2026-2027 recruitment cycle in Hanoi. With five years of dedicated IELTS instruction experience, official IELTS examiner credentials, and a demonstrated track record of helping students achieve band scores of 7.0 and above, I am confident I would make a valuable addition to your academic team.

In my current role as IELTS Teacher at RMIT English School (2018-2024), I have taught across all four papers -- Listening, Reading, Writing, and Speaking -- with particular expertise in the speaking and writing components that present the greatest challenge to candidates. My students have consistently achieved band scores of 7.0 and above, and I have developed structured, evidence-based revision programmes that target the specific weaknesses typical of Vietnamese learners at the 5.5-6.5 band range. I am experienced working with secondary and upper-secondary students aged 15-18, which aligns directly with your in-school programme.

My PGCE studies through Liverpool John Moores University (2024-2026) are strengthening my pedagogical practice with a particular focus on secondary education, curriculum design, and differentiated instruction. I hold a Bachelor of Business Management & Finance, TEFL/TESOL certification, and professional certifications in Project Management (CAPM(R)) and Agile Software Development (DASM), all of which enrich my approach to language teaching and student assessment.

I am fully conversant with modern educational technology platforms and have experience teaching Cambridge, IB, and national curricula. I am excited about Atlantic Five-Star English's in-school model and the opportunity to contribute to a structured, long-term programme starting September 2026.

Thank you for your consideration. I look forward to hearing from you.

Sincerely,
{PROFILE['name']}
{PROFILE['email']} | {PROFILE['phone']}
"""

# Save updated IELTS cover letter
out_path = '/home/amrba/job_scraper/vietnam/cover_letters/Atlantic Teacher Careers_IELTS Teacher_2026-2027_Hanoi_20260513.docx'
doc = Document()
for para in content.split('\n\n'):
    if para.strip():
        p = doc.add_paragraph()
        p.add_run(para.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(8)
doc.save(out_path)
print(f"Saved: {out_path}")
